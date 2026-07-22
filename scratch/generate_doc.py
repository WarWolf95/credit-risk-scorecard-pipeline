import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_heading_styled(doc, text, level, space_before=12, space_after=6):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(180, 40, 40)  # Alert Crimson/Deep Red for Hyper-plan
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(89, 89, 89)
        run.bold = True
    return p

def create_hyper_plan():
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # Title Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("MAULIK'S 3-WEEK HYPER-ACCELERATION PLAN")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(180, 40, 40)
    title_run.bold = True
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    sub_run = subtitle_p.add_run("Zero-Fluff, High-Intensity Path to UK Job Readiness (10-12 Hours/Day)")
    sub_run.font.name = 'Segoe UI'
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(127, 127, 127)
    sub_run.font.italic = True

    # Callout Box / Rules of the Sprint
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "FFF5F5")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24') # 3pt width
    left_border.set(qn('w:space'), '0')
    left_border.set(qn('w:color'), 'B42828')
    tcBorders.append(left_border)
    for side in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(4)
    r1 = cp.add_run("CRITICAL SPEED PROTOCOL\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(180, 40, 40)
    
    r2 = cp.add_run(
        "To compress an 8-week timeline into 21 days, Maulik must bypass academic theory, ignore secondary toolkits, "
        "and focus entirely on active retrieval and output generation. He will build exactly ONE world-class flagship "
        "portfolio project (The Credit Risk Scorecard model), master window functions, lock in relational data modeling, "
        "and aggressively rewrite his professional profile for the UK job market."
    )
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 1: The 21-Day Day-by-Day Sprint Schedule
    add_heading_styled(doc, "1. The 21-Day Day-by-Day Sprint Schedule", level=1)
    
    plan_table = doc.add_table(rows=4, cols=3)
    plan_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(plan_table, color="D3D3D3")
    
    headers_plan = ["Timeframe", "Daily Action Items & Concepts to Lock In", "Key Deliverables / Proof of Competence"]
    hdr_cells_plan = plan_table.rows[0].cells
    for i, title in enumerate(headers_plan):
        hdr_cells_plan[i].text = title
        set_cell_background(hdr_cells_plan[i], "B42828")
        set_cell_margins(hdr_cells_plan[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells_plan[i].paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    plan_data = [
        ("Week 1 (Days 1 - 7)\nSQL Mastery & BI Dashboards",
         "• Days 1-3: Extreme SQL drill. Focus solely on Window Functions (RANK, LEAD/LAG, SUM/AVG OVER), CTEs, and self-joins. Work through Mode Analytics SQL Advanced module + solve 40 medium problems on StrataScratch.\n"
         "• Days 4-5: Power BI acceleration. Build a relational star-schema model. Write complex DAX measures (CALCULATE, FILTER, context transitions). Avoid simple flat tables.\n"
         "• Days 6-7: Practical Stats. Master A/B testing (hypothesis setups, T-test coding, statistical significance, sample size requirements).",
         "🏆 SQL interview-readiness verified (StrataScratch dashboard screens solved).\n"
         "🏆 Clean Power BI Dashboard connected to a relational database, published/screenshot-ready."),
         
        ("Week 2 (Days 8 - 14)\nFlagship Project & API Build",
         "• Days 8-9: Pandas vectorisation & data cleaning drills. Master groupby, merges, and pivot tables without looking at documentation.\n"
         "• Days 10-12: Flagship Credit Risk Scorecard. Take the credit dataset, execute feature engineering (WoE/IV encoding), train a Logistic Regression model, scale to scorecard points. Train an alternative XGBoost/LightGBM model and compare results using ROC-AUC, Precision, and Recall metrics.\n"
         "• Days 13-14: API & Local UI. Wrap the model using FastAPI, build a basic Streamlit UI showing model output comparison, and commit code with a professional README to GitHub.",
         "🏆 Flagship Repository live on GitHub: complete credit risk pipeline (Python code, configuration parameters, and detailed markdown explanation).\n"
         "🏆 Streamlit web application running locally/deployed showing interactive credit scorecard calculations."),
         
        ("Week 3 (Days 15 - 21)\nProfile Re-brand & Apply Sprint",
         "• Days 15-16: Rewrite CV & LinkedIn. Reframe Numerator India operations. Rebrand 'training legacy software' as 'Operations Quality Assurance and Data Standards lead'. Emphasize business value, metrics, and automation.\n"
         "• Days 17-18: Intensive live interview mock drills. Practice talking through the Credit Risk Scorecard project decisions (e.g., why choose Logistic Regression, how WoE works, why optimize Recall).\n"
         "• Days 19-21: High-volume application & networking blitz. Apply to 15-20 junior/mid analyst and data science roles daily. Message 10 UK engineering/analytics managers per day with a direct, short link to his flagship project.",
         "🏆 Optimised CV & LinkedIn profiles live.\n"
         "🏆 50+ targeted applications submitted + 30 direct networking connections established with UK managers.")
    ]
    
    for row_idx, data in enumerate(plan_data):
        row_cells = plan_table.rows[row_idx + 1].cells
        bg_color = "FFFBFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(9.0)
            if col_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(180, 40, 40)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 2: Zero-Fluff Tactical Focus
    add_heading_styled(doc, "2. Tactical Shortcuts (Skip vs. Double Down)", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "With a compressed timeline, Maulik cannot afford to study everything. Here is exactly what he must skip, "
        "and where he must double down to maximize UK interview conversion."
    )
    p.paragraph_format.space_after = Pt(8)
    
    t_tbl = doc.add_table(rows=6, cols=3)
    t_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_tbl, color="D3D3D3")
    
    headers_t = ["Concept Area", "❌ SKIP ENTIRELY (Save 20+ Hours)", "🔥 DOUBLE DOWN (Master Deeply)"]
    hdr_cells_t = t_tbl.rows[0].cells
    for i, title in enumerate(headers_t):
        hdr_cells_t[i].text = title
        set_cell_background(hdr_cells_t[i], "555555")
        set_cell_margins(hdr_cells_t[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells_t[i].paragraphs[0]
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    t_data = [
        ("SQL", "Complex indexing structures, database admin rules, database migration scripts.", "Window functions, self-joins, CTEs, query readability."),
        ("Python & Wrangle", "Advanced OOP, package development, decorator building, async python.", "Pandas vectorisation, groupby aggregates, handling null values."),
        ("Statistics", "Bayesian math derivations, ANOVA tables, non-parametric proof details.", "A/B testing, hypothesis testing, T-tests, p-values."),
        ("Machine Learning", "Neural networks, Deep Learning, NLP transformers, Computer Vision.", "Logistic Regression, Random Forests, XGBoost, Model metrics (ROC-AUC)."),
        ("BI & Visualisation", "Learning 3 different BI tools (Tableau, Power BI, Qlik). Choose one.", "Power BI: Relational star-schema design, DAX measures.")
    ]
    
    for row_idx, data in enumerate(t_data):
        row_cells = t_tbl.rows[row_idx + 1].cells
        bg_color = "F9F9F9" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(9.0)
            if col_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(51, 51, 51)
            elif col_idx == 1:
                run.font.color.rgb = RGBColor(150, 50, 50)
            elif col_idx == 2:
                run.font.color.rgb = RGBColor(30, 100, 30)
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: The 3-Step Re-branding Strategy
    add_heading_styled(doc, "3. The 3-Step Profile Re-branding Strategy", level=1)
    
    p = doc.add_paragraph()
    p.add_run(
        "Maulik's 6.5 years in Numerator India is his secret weapon. Fresh MSc grads lack corporate maturity. "
        "He must present his operations background as analytical leadership:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    reb_points = [
        ("Step 1: Reframe legacy QC as 'Data Quality'", "Instead of 'Checked accuracy of inputs,' write: 'Established data validation protocols and anomaly detection checks to maintain data integrity across core reporting pipelines.'"),
        ("Step 2: Reframe reporting as 'BI Development'", "Instead of 'Created performance and accuracy reports,' write: 'Designed and deployed analytical dashboards using Power BI to visualize operational metrics, improving productivity and efficiency.'"),
        ("Step 3: Reframe learning as 'Process Automation'", "Instead of 'Learned Python/SQL,' write: 'Identified workflow bottlenecks and proactively wrote custom SQL scripts and Python automation to replace manual data processing tasks, reducing turnaround times.'")
    ]
    for title, desc in reb_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_title = p.add_run(f"{title}: ")
        r_title.bold = True
        p.add_run(desc)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 4: Highly Effective Learning Resources
    add_heading_styled(doc, "4. High-Impact Learning Resources (Tested & Curated)", level=1)
    
    res_intro = doc.add_paragraph()
    res_intro.add_run("Use these specific, interactive sources to construct active learning loops. Avoid passive watching.")
    res_intro.paragraph_format.space_after = Pt(8)
    
    resources_list = [
        ("SQL & Database Mastery", 
         "1. Mode Analytics SQL Tutorial (Intermediate & Advanced Modules): Interactive in-browser SQL queries using real datasets. Crucial for CTEs, Window Functions, and basic DB modeling.\n"
         "2. Alex The Analyst SQL Bootcamp (YouTube): Excellent for learning visual query structures, Temp Tables, and cleaning real-world datasets step-by-step."),
         
        ("Python & Wrangling",
         "1. Kaggle Learn (Pandas & Python Modules): Free, high-density, notebook-based coding tutorials. Excellent for mastering data aggregation and cleaning fast.\n"
         "2. StrataScratch Python Practice: Switch the query parser to Python/Pandas mode and solve SQL equivalent problems using DataFrame manipulations."),
         
        ("Statistics & A/B Testing",
         "1. StatQuest with Josh Starmer (YouTube): Visual breakdowns of distributions, p-values, hypothesis tests, and basic regression math without confusing symbols.\n"
         "2. Udacity's Free A/B Testing Course (by Google): The industry gold standard for learning metrics, statistical significance, and operational A/B design."),
         
        ("Machine Learning & Visualisation",
         "1. Kaggle Intro & Intermediate Machine Learning: Quick, hands-on tutorials using Random Forests and XGBoost, with a focus on metrics and pipeline construction.\n"
         "2. SQLBI / sqlbi.com (Alberto Ferrari & Marco Russo): The absolute best resource to understand Power BI DAX context transitions, star-schema modeling, and dashboard modeling principles.")
    ]
    
    for title, details in resources_list:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(f"{title}\n")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(180, 40, 40)
        r_title.font.size = Pt(11)
        
        p_desc = doc.add_paragraph()
        p_desc.add_run(details)
        p_desc.paragraph_format.space_after = Pt(10)

    # Save the document
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    doc_path = os.path.join(desktop_path, "Maulik_3Week_Hyper_Acceleration_Plan.docx")
    doc.save(doc_path)
    print(f"Hyper-plan successfully saved to: {doc_path}")

if __name__ == "__main__":
    create_hyper_plan()
